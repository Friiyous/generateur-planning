"""
Générateur de Planning - Centre de Santé
Application Streamlit pour générer automatiquement les plannings du personnel de santé.
Cycle: PG → R → P
"""

import streamlit as st
import pandas as pd
from datetime import datetime
import os

# Configuration de la page
st.set_page_config(
    page_title="Générateur de Planning",
    page_icon="🏥",
    layout="wide"
)

# Fichier de données
DATA_FILE = "employes.json"

# Catégories
CATEGORIES = [
    {"id": "infirmier-dispensaire", "label": "Infirmiers (Dispensaire)", "prefixe": "Infirmier"},
    {"id": "aide-dispensaire", "label": "Aides (Dispensaire)", "prefixe": "Aide"},
    {"id": "sage-femme-maternite", "label": "Sages-femmes (Maternité)", "prefixe": "Sage-femme"},
    {"id": "aide-maternite", "label": "Aides (Maternité)", "prefixe": "Aide"},
    {"id": "fille-salle", "label": "Filles de salle", "prefixe": "Fille de salle"},
]

# Cycle: PG → R → P
CYCLE = ["PG", "R", "P"]

def charger_employes():
    """Charger les employés depuis le fichier JSON"""
    try:
        if os.path.exists(DATA_FILE):
            import json
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []

def sauvegarder_employes(employes):
    """Sauvegarder les employés dans le fichier JSON"""
    import json
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(employes, f, ensure_ascii=False, indent=2)

def generer_planning_employe(cycle_position, annee, mois):
    """Générer le planning pour un employé"""
    import calendar
    jours_dans_mois = calendar.monthrange(annee, mois + 1)[1]
    planning = []

    position = cycle_position
    for jour in range(1, jours_dans_mois + 1):
        planning.append({
            "jour": f"{jour:02d}",
            "shift": CYCLE[position]
        })
        position = (position + 1) % len(CYCLE)

    return planning, position

def creer_en_tete_officiel(centre_sante=""):
    """Créer l'en-tête officiel pour Word"""
    en_tete = [
        "MINISTÈRE DE LA SANTÉ, DE L'HYGIÈNE PUBLIQUE ET DE LA COUVERTURE MALADIE UNIVERSELLE",
        "RÉPUBLIQUE DE CÔTE D'IVOIRE",
        "Union - Discipline - Travail",
        "_______________________________________",
    ]
    # Ajouter le nom du centre si fourni
    if centre_sante:
        en_tete.append(f"Centre: {centre_sante}")
        en_tete.append("")
    else:
        en_tete.append("")
    return en_tete

def exporter_word(plannings_data, service_label, mois, annee, centre_sante=""):
    """Exporter le planning en Word"""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    doc = Document()

    # Format PAYSAGE
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # En-tête officiel
    for ligne in creer_en_tete_officiel(centre_sante):
        p = doc.add_paragraph(ligne)
        if "MINISTÈRE" in ligne:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif "RÉPUBLIQUE" in ligne or "Union" in ligne:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Titre du planning
    titre = doc.add_paragraph()
    titre.add_run(f"PLANNING MENSUEL - {mois} {annee}").bold = True
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Service
    service_p = doc.add_paragraph(f"Service: {service_label}")
    if centre_sante:
        centre_p = doc.add_paragraph(f"Centre: {centre_sante}")
    service_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Fonction pour appliquer couleur de fond
    def set_shading(cell, color_hex):
        """Appliquer couleur de fond avec shading"""
        from docx.oxml import OxmlElement
        tc = cell._element.get_or_add_tcPr()
        # Supprimer ancien shading
        for shd in tc.findall(qn('w:shd')):
            tc.remove(shd)
        # Ajouter nouveau shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tc.append(shd)

    # Créer le tableau - Sans couleur (juste en-tête bleu)
    if plannings_data:
        jours = len(plannings_data[0]["planning"])
        table = doc.add_table(rows=len(plannings_data) + 1, cols=jours + 1)
        table.style = "Table Grid"

        # En-tête - BLEU seulement
        header_cells = table.rows[0].cells
        header_cells[0].text = "NOM & PRENOM"
        set_shading(header_cells[0], "3498DB")  # Bleu

        for i, j in enumerate(range(jours)):
            header_cells[i + 1].text = f"{j + 1:02d}"
            set_shading(header_cells[i + 1], "3498DB")  # Bleu

        # Données - TOUT SANS COULEUR (fond blanc, texte noir)
        for i, emp in enumerate(plannings_data):
            row = table.rows[i + 1].cells
            row[0].text = f"{emp['nom']} {emp['prenom']}"
            # Pas de couleur pour Nom

            for j, shift in enumerate(emp["planning"]):
                cell = row[j + 1]
                cell.text = shift["shift"]
                # PAS de couleur - tout reste en blanc avec texte noir

    return doc

# TOUJOURS recharger depuis le fichier directement
employes_fresh = charger_employes()

# Titre
st.title("🏥 Générateur de Planning - Centre de Santé")
st.markdown("### Cycle automatique: PG → R → P")

# Légende
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown("**PG** = Permanence + Garde")
with col2:
    st.markdown("**P** = Permanence")
with col3:
    st.markdown("**R** = Repos")

st.divider()

# Gestion du personnel par catégorie
st.header("👥 Gestion du Personnel")

# Créer des colonnes pour les catégories
cols = st.columns(5)

for idx, cat in enumerate(CATEGORIES):
    with cols[idx]:
        st.subheader(cat["label"])

        # Filtrer les employés de cette catégorie
        employes_cat = [e for e in employes_fresh if e["service"] == cat["id"]]

        # Afficher le nombre
        st.caption(f"{len(employes_cat)} employé(s)")

        # Formulaire pour ajouter
        with st.form(f"form_{cat['id']}"):
            nouveau_nom = st.text_input(f"Nom", placeholder="Ex: YEO", key=f"nom_{cat['id']}")
            submitted = st.form_submit_button("➕ Ajouter", use_container_width=True)

            if submitted and nouveau_nom:
                # Trouver la prochaine position de cycle décalée
                employes_meme_service = [e for e in employes_fresh if e["service"] == cat["id"]]
                if employes_meme_service:
                    # Prendre la position suivante dans le cycle
                    positions = [e["cyclePosition"] for e in employes_meme_service]
                    nouvelle_pos = (max(positions) + 1) % len(CYCLE)
                else:
                    nouvelle_pos = 0

                nouveau = {
                    "id": len(employes_fresh) + 1,
                    "nom": nouveau_nom.upper(),
                    "prenom": cat["prefixe"],
                    "service": cat["id"],
                    "cyclePosition": nouvelle_pos
                }
                employes_fresh.append(nouveau)
                sauvegarder_employes(employes_fresh)
                st.rerun()

        # Liste des employés
        for emp in employes_cat:
            col_a, col_b = st.columns([3, 1])
            with col_a:
                st.markdown(f"**{emp['nom']}**")
            with col_b:
                if st.button("×", key=f"del_{emp['id']}_{cat['id']}"):
                    employes_fresh = [e for e in employes_fresh if e["id"] != emp["id"]]
                    sauvegarder_employes(employes_fresh)
                    st.rerun()

st.divider()

# Générateur de planning
st.header("📅 Générateur de Planning")

# Nom du centre de santé
centre_sante = st.text_input("🏥 Nom du Centre de Santé", placeholder="Ex: Centre de Santé de Korhogo")

# Options de génération
col1, col2, col3 = st.columns(3)

with col1:
    mois = st.selectbox("Mois", [
        "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
        "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"
    ], index=datetime.now().month - 1)

with col2:
    annee = st.number_input("Année", min_value=2020, max_value=2030, value=datetime.now().year)

with col3:
    # Option pour choisir un service spécifique ou tous
    generer_tous = st.checkbox("Générer pour tous les services", value=False)
    if generer_tous:
        service = None
    else:
        service = st.selectbox("Service", CATEGORIES, format_func=lambda x: x["label"])

if st.button("🔄 Générer le Planning", type="primary", use_container_width=True):
    # Filtrer les employés du service
    if service is None:
        # Générer pour tous les services
        employes_service = employes_fresh
    else:
        employes_service = [e for e in employes_fresh if e["service"] == service["id"]]

    if not employes_service:
        st.warning("Aucun employé dans ce service ! Ajoutez d'abord des employés.")
    else:
        mois_num = ["Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
                 "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"].index(mois)

        # Recharger les positions depuis le fichier AVANT de générer
        employes_fresh = charger_employes()
        employes_service = [e for e in employes_fresh if e["service"] == service["id"]]

        plannings = []
        for emp in employes_service:
            planning, nouvelle_position = generer_planning_employe(emp["cyclePosition"], annee, mois_num)
            plannings.append({
                "nom": emp["nom"],
                "prenom": emp["prenom"],
                "planning": planning
            })

        # NE PAS sauvegarder les positions - les garder fixes!

        # Afficher le planning
        st.success(f"✅ Planning généré pour {len(plannings)} employé(s)")

        # Créer un tableau HTML stylisé
        jours = len(plannings[0]["planning"])

        # En-tête du tableau HTML
        html = """
        <style>
            .planning-table { border-collapse: collapse; width: 100%; font-size: 11px; }
            .planning-table th, .planning-table td { border: 1px solid #fff; padding: 3px; text-align: center; }
            .planning-table th { background: #3498db; color: white; font-weight: bold; }
            .planning-table th:first-child { background: #2c3e50; }
            .planning-table td:first-child { background: #ecf0f1; font-weight: bold; text-align: left; color: #2c3e50; }
            .shift-PG { background: #e74c3c; color: white; font-weight: bold; padding: 2px 4px; border-radius: 2px; }
            .shift-P { background: #3498db; color: white; font-weight: bold; padding: 2px 4px; border-radius: 2px; }
            .shift-R { background: #95a5a6; color: white; font-weight: bold; padding: 2px 4px; border-radius: 2px; }
        </style>
        <div style="overflow-x: auto;">
        <table class="planning-table">
            <thead>
                <tr>
                    <th>Nom & Prénom</th>
        """

        # Ajouter les en-têtes de jours
        for i in range(jours):
            html += f'<th>{i+1:02d}</th>'

        html += '</tr></thead><tbody>'

        # Ajouter les lignes d'employés
        for emp in plannings:
            html += f"<tr><td>{emp['nom']} {emp['prenom']}</td>"
            for shift in emp["planning"]:
                html += f'<td><span class="shift-{shift["shift"]}">{shift["shift"]}</span></td>'
            html += '</tr>'

        html += '</tbody></table></div>'

        st.markdown(html, unsafe_allow_html=True)

        # Bouton export Word
        service_label = service["label"] if service else "Tous les services"
        doc = exporter_word(plannings, service_label, mois, annee, centre_sante)

        # Sauvegarder temporairement
        service_id = service["id"] if service else "tous"
        temp_file = f"planning_{service_id}_{mois}_{annee}.docx"
        doc.save(temp_file)

        with open(temp_file, "rb") as f:
            st.download_button(
                "📄 Exporter en Word",
                f.read(),
                file_name=temp_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.divider()

# Total des employés
total = len(employes_fresh)
st.caption(f"Total: {total} employé(s) enregistré(s)")